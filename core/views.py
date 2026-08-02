from io import BytesIO
import hashlib
import json
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import Group,User
from django.contrib.auth.views import LoginView
from django.db import transaction
from django.db.models import Q
from django.http import HttpResponse
from django.shortcuts import get_object_or_404,redirect,render
from django.urls import reverse
from django.views.decorators.http import require_POST
from django.utils import timezone
from .models import *
from .forms import *
from .services import project_summary,trace_rows
from .permissions import ROLE_LABELS,admin_required,can_modify,has_full_access,user_role

MODEL_MAP={'requirements':(Requirement,RequirementForm,'요구사항'),'designs':(DesignItem,DesignItemForm,'설계'),'risks':(Risk,RiskForm,'위험'),'tests':(TestCase,TestCaseForm,'시험'),'incidents':(Incident,IncidentForm,'이상사례'),'capas':(CAPA,CAPAForm,'CAPA'),'changes':(ChangeRequest,ChangeRequestForm,'변경 요청')}
def can_edit(user): return can_modify(user)
def list_item(obj,kind):
    """서로 다른 모델을 목록 템플릿용 공통 표현으로 변환한다."""
    title=getattr(obj,'name',None) or getattr(obj,'title',None) or getattr(obj,'hazard',None) or getattr(obj,'test_name',None) or getattr(obj,'incident_summary',None) or '-'
    status=getattr(obj,'status',None) or getattr(obj,'investigation_status',None) or '-'
    code=getattr(obj,'code',None) or getattr(obj,'project_id',None) or f'ID-{obj.pk}'
    edit_url=reverse('project_edit',args=[obj.pk]) if kind=='projects' else reverse('object_edit',args=[kind,obj.pk])
    delete_url=None if kind=='projects' else reverse('object_delete',args=[kind,obj.pk])
    locked=kind=='changes' and hasattr(obj,'approval_signature')
    return {'pk':obj.pk,'code':code,'title':title,'status':status,'updated_at':getattr(obj,'updated_at',None),'edit_url':edit_url,'delete_url':delete_url,'has_result':kind=='tests','locked':locked,'approve_url':reverse('change_approve',args=[obj.pk]) if kind=='changes' and not locked else None,'revoke_url':reverse('change_revoke',args=[obj.pk]) if kind=='changes' and locked else None}
@login_required
def dashboard(request):
    project=Project.objects.first(); summary=project_summary(project) if project else None
    unread=Notification.objects.filter(user=request.user,is_read=False).count()
    return render(request,'dashboard.html',{'project':project,'summary':summary,'projects':Project.objects.all(),'recent':AuditLog.objects.order_by('-created_at')[:8],'unread_notifications':unread})
@login_required
def projects(request):
    return render(request,'list.html',{'title':'프로젝트','items':[list_item(x,'projects') for x in Project.objects.all()],'kind':'projects'})
@login_required
def project_form(request,pk=None):
    if not can_edit(request.user): return HttpResponse('수정 권한이 없습니다.',status=403)
    obj=get_object_or_404(Project,pk=pk) if pk else None; form=ProjectForm(request.POST or None,instance=obj)
    if request.method=='POST' and form.is_valid():
        x=form.save(commit=False); x.created_by=x.created_by or request.user; x.save(); messages.success(request,'프로젝트가 저장되었습니다.'); return redirect('projects')
    return render(request,'form.html',{'form':form,'title':'프로젝트 저장'})
@login_required
def object_list(request,kind):
    model,form,label=MODEL_MAP[kind]; project=get_object_or_404(Project,pk=request.GET.get('project') or Project.objects.values_list('pk',flat=True).first()); qs=model.objects.filter(project=project)
    q=request.GET.get('q','').strip(); status=request.GET.get('status','').strip()
    if q:
        query=Q(code__icontains=q)
        for field in ('title','name','hazard','description','reason','impact'):
            if field in {f.name for f in model._meta.fields}: query |= Q(**{f'{field}__icontains':q})
        qs=qs.filter(query)
    if status and 'status' in {f.name for f in model._meta.fields}: qs=qs.filter(status=status)
    statuses=model.objects.filter(project=project).exclude(status='').values_list('status',flat=True).distinct() if 'status' in {f.name for f in model._meta.fields} else []
    return render(request,'list.html',{'title':label,'items':[list_item(x,kind) for x in qs],'kind':kind,'project':project,'statuses':statuses,'q':q,'selected_status':status})
@login_required
def object_form(request,kind,pk=None):
    if not can_edit(request.user): return HttpResponse('수정 권한이 없습니다.',status=403)
    model,form_cls,label=MODEL_MAP[kind]; obj=get_object_or_404(model,pk=pk) if pk else None
    if kind=='changes' and obj and hasattr(obj,'approval_signature'): return HttpResponse('전자서명으로 승인된 변경 요청은 수정할 수 없습니다.',status=409)
    original_data={field.name:str(getattr(obj,field.name,'')) for field in obj._meta.fields} if obj else {}
    project=obj.project if obj else get_object_or_404(Project,pk=request.GET.get('project') or Project.objects.values_list('pk',flat=True).first()); form=form_cls(request.POST or None,instance=obj)
    if request.method=='POST' and form.is_valid():
        before=original_data
        if obj:
            VersionSnapshot.objects.create(project=project,model_name=model.__name__,object_id=obj.pk,object_code=getattr(obj,'code',''),version=VersionSnapshot.objects.filter(model_name=model.__name__,object_id=obj.pk).count()+1,data=before,change_reason=request.POST.get('change_reason',''),changed_by=request.user)
        x=form.save(commit=False); x.project=project; x.created_by=x.created_by or request.user
        if kind=='changes' and x.status in ('승인','반려') and (not obj or obj.status != x.status):
            x.reviewed_by=request.user; x.reviewed_at=timezone.now()
        x.save(); form.save_m2m()
        after={field.name:str(getattr(x,field.name,'')) for field in x._meta.fields}
        changed={key:{'before':before.get(key),'after':value} for key,value in after.items() if before.get(key)!=value}
        AuditLog.objects.create(user=request.user,action='생성' if not obj else '수정',path=request.path[:300],method=request.method,ip=request.META.get('REMOTE_ADDR'),details={'model':model.__name__,'object_id':x.pk,'code':getattr(x,'code',''),'changes':changed},response_status=302)
        if kind=='changes':
            recipients={u for u in (x.requester,x.assignee,project.manager) if u and u != request.user}
            for user in recipients:
                Notification.objects.create(user=user,project=project,kind='변경 요청',title=f'{x.code} {x.status}',message=x.title,link=reverse('object_edit',args=['changes',x.pk]))
        messages.success(request,f'{label} 항목이 저장되었습니다.'); return redirect('object_list',kind=kind)
    return render(request,'form.html',{'form':form,'title':f'{label} 저장','project':project})
@require_POST
@login_required
def object_delete(request,kind,pk):
    if not can_edit(request.user): return HttpResponse(status=403)
    if kind=='changes':
        obj=get_object_or_404(ChangeRequest,pk=pk)
        if hasattr(obj,'approval_signature'): return HttpResponse('전자서명으로 승인된 변경 요청은 삭제할 수 없습니다.',status=409)
    MODEL_MAP[kind][0].objects.filter(pk=pk).delete(); messages.success(request,'삭제되었습니다.'); return redirect('object_list',kind=kind)
@login_required
def result_form(request,pk):
    test=get_object_or_404(TestCase,pk=pk); result=getattr(test,'result',None); form=TestResultForm(request.POST or None,instance=result)
    if request.method=='POST' and form.is_valid():
        x=form.save(commit=False); x.test_case=test; x.created_by=request.user; x.save(); return redirect('object_list',kind='tests')
    return render(request,'form.html',{'form':form,'title':f'{test.code} 결과 등록'})
@login_required
def traceability(request):
    project=get_object_or_404(Project,pk=request.GET.get('project') or Project.objects.values_list('pk',flat=True).first()); rows=trace_rows(project)
    if request.GET.get('gaps')=='1': rows=[r for r in rows if r['status']!='완료']
    return render(request,'traceability.html',{'project':project,'rows':rows})
@login_required
def export(request,fmt):
    project=get_object_or_404(Project,pk=request.GET.get('project') or Project.objects.values_list('pk',flat=True).first()); rows=trace_rows(project)
    if fmt=='xlsx':
        from openpyxl import Workbook
        wb=Workbook(); ws=wb.active; ws.title='Traceability'; ws.append(['요구사항','제목','설계','위험','시험','결과','CAPA','상태'])
        for r in rows: ws.append([r['req'].code,r['req'].title,','.join(x.code for x in r['designs']),','.join(x.code for x in r['risks']),','.join(x.code for x in r['tests']),','.join(getattr(getattr(x,'result',None),'outcome','미등록') for x in r['tests']),','.join(x.code for x in r['capas']),r['status']])
        b=BytesIO(); wb.save(b); data=b.getvalue(); content='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    elif fmt=='docx':
        from docx import Document
        d=Document(); d.add_heading(f'{project.name} 추적성 보고서',0); d.add_paragraph('포트폴리오·교육용 초안 — 담당자 검토 필요')
        for r in rows: d.add_heading(f"{r['req'].code} {r['req'].title}",1); d.add_paragraph(f"상태: {r['status']} / 위험: {', '.join(x.code for x in r['risks']) or '-'} / 시험: {', '.join(x.code for x in r['tests']) or '-'}")
        b=BytesIO(); d.save(b); data=b.getvalue(); content='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        from reportlab.pdfgen import canvas
        b=BytesIO(); c=canvas.Canvas(b); c.drawString(40,800,'MedTrace RA - Traceability Report'); y=775
        for r in rows: c.drawString(40,y,f"{r['req'].code} | {r['status']}"); y-=18
        c.save(); data=b.getvalue(); content='application/pdf'
    response=HttpResponse(data,content_type=content); response['Content-Disposition']=f'attachment; filename="traceability.{fmt}"'; return response
@admin_required
def audit(request):
    return render(request,'audit.html',{'items':AuditLog.objects.order_by('-created_at')[:200]})

@login_required
def notifications(request):
    items=Notification.objects.filter(user=request.user)
    if request.GET.get('read')=='all':
        items.filter(is_read=False).update(is_read=True)
        return redirect('notifications')
    return render(request,'notifications.html',{'items':items})

@login_required
def requirement_import(request):
    if not can_edit(request.user): return HttpResponse('수정 권한이 없습니다.',status=403)
    form=RequirementImportForm(request.POST or None,request.FILES or None)
    if request.method=='POST' and form.is_valid():
        from openpyxl import load_workbook
        project=form.cleaned_data['project']
        try:
            ws=load_workbook(form.cleaned_data['file'],read_only=True,data_only=True).active
            headers=[str(x or '').strip().lower() for x in next(ws.iter_rows(values_only=True))]
            required={'title','description'}
            if not required.issubset(headers): raise ValueError('title, description 열이 필요합니다.')
            created=0
            with transaction.atomic():
                for row in ws.iter_rows(values_only=True):
                    data=dict(zip(headers,row))
                    if not data.get('title'): continue
                    Requirement.objects.create(project=project,code='',title=str(data['title']),description=str(data.get('description') or ''),req_type=str(data.get('req_type') or '기능 요구사항'),priority=str(data.get('priority') or '중'),status=str(data.get('status') or '초안'),created_by=request.user)
                    created+=1
            messages.success(request,f'요구사항 {created}건을 가져왔습니다.')
            return redirect(f"{reverse('object_list',args=['requirements'])}?project={project.pk}")
        except Exception as exc:
            form.add_error('file',f'가져오기에 실패했습니다: {exc}')
    return render(request,'form.html',{'form':form,'title':'요구사항 Excel 가져오기'})

@admin_required
def user_management(request,pk=None):
    target=get_object_or_404(User,pk=pk) if pk else None
    initial=None
    if target:
        initial={'username':target.username,'email':target.email,'role':user_role(target),'is_active':target.is_active}
    form=ManagedUserForm(request.POST or None,user_instance=target,initial=initial)
    if request.method=='POST' and form.is_valid():
        if target==request.user and (form.cleaned_data['role']!='ADMIN' or not form.cleaned_data['is_active']):
            form.add_error('role','현재 로그인한 관리자 계정은 관리자 권한을 해제하거나 비활성화할 수 없습니다.')
        else:
            user=target or User()
            user.username=form.cleaned_data['username']; user.email=form.cleaned_data['email']; user.is_active=form.cleaned_data['is_active']
            role=form.cleaned_data['role']; user.is_staff=role=='ADMIN'; user.is_superuser=role=='ADMIN'
            if form.cleaned_data['password']: user.set_password(form.cleaned_data['password'])
            user.save()
            user.groups.clear(); group,_=Group.objects.get_or_create(name=role); user.groups.add(group)
            messages.success(request,f'{user.username} 계정을 {ROLE_LABELS[role]} 역할로 저장했습니다.')
            return redirect('user_management')
    users=[{'pk':u.pk,'username':u.username,'email':u.email,'role':ROLE_LABELS.get(user_role(u),user_role(u)),'active':u.is_active,'last_login':u.last_login} for u in User.objects.order_by('username')]
    return render(request,'users.html',{'items':users,'form':form,'target':target})

@login_required
def change_approve(request,pk):
    if not can_edit(request.user): return HttpResponse('승인 권한이 없습니다.',status=403)
    change=get_object_or_404(ChangeRequest,pk=pk)
    if hasattr(change,'approval_signature'): return HttpResponse('이미 승인된 변경 요청입니다.',status=409)
    form=ApprovalSignatureForm(request.POST or None)
    if request.method=='POST' and form.is_valid():
        if not request.user.check_password(form.cleaned_data['password']):
            form.add_error('password','비밀번호가 일치하지 않습니다.')
        else:
            payload={field.name:str(getattr(change,field.name,'')) for field in change._meta.fields}
            payload['requirements']=list(change.requirements.order_by('pk').values_list('pk',flat=True))
            digest=hashlib.sha256(json.dumps(payload,sort_keys=True,ensure_ascii=False).encode('utf-8')).hexdigest()
            with transaction.atomic():
                VersionSnapshot.objects.create(project=change.project,model_name='ChangeRequest',object_id=change.pk,object_code=change.code,version=VersionSnapshot.objects.filter(model_name='ChangeRequest',object_id=change.pk).count()+1,data=payload,change_reason='전자서명 승인',changed_by=request.user)
                ApprovalSignature.objects.create(change_request=change,signer=request.user,comment=form.cleaned_data['comment'],content_hash=digest)
                ApprovalEvent.objects.create(change_request=change,event='approved',actor=request.user,comment=form.cleaned_data['comment'],content_hash=digest)
                change.status='승인'; change.reviewed_by=request.user; change.reviewed_at=timezone.now(); change.save(update_fields=['status','reviewed_by','reviewed_at','updated_at'])
            messages.success(request,f'{change.code} 변경 요청을 전자서명으로 승인했습니다.')
            return redirect('object_list',kind='changes')
    return render(request,'approval.html',{'form':form,'change':change,'title':f'{change.code} 전자서명 승인'})

@login_required
def change_revoke(request,pk):
    if not can_edit(request.user): return HttpResponse('승인 취소 권한이 없습니다.',status=403)
    change=get_object_or_404(ChangeRequest,pk=pk)
    signature=getattr(change,'approval_signature',None)
    if not signature: return HttpResponse('활성화된 승인이 없습니다.',status=409)
    form=ApprovalRevokeForm(request.POST or None)
    if request.method=='POST' and form.is_valid():
        if not request.user.check_password(form.cleaned_data['password']):
            form.add_error('password','비밀번호가 일치하지 않습니다.')
        else:
            with transaction.atomic():
                ApprovalEvent.objects.create(change_request=change,event='revoked',actor=request.user,comment=form.cleaned_data['reason'],content_hash=signature.content_hash)
                signature.delete()
                change.status='검토 중'; change.reviewed_by=None; change.reviewed_at=None
                change.save(update_fields=['status','reviewed_by','reviewed_at','updated_at'])
            messages.success(request,f'{change.code} 승인을 취소했습니다. 변경 후 재승인이 필요합니다.')
            return redirect('object_list',kind='changes')
    return render(request,'approval.html',{'form':form,'change':change,'title':f'{change.code} 승인 취소'})
